#!/usr/bin/env python3
"""
Professional Resume Word Generator
Creates ATS-friendly Word documents using python-docx with precise formatting control
"""

import argparse
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class ResumeData:
    """Structured resume data (same as PDF generator)"""
    
    def __init__(self):
        self.personal_info = {
            'name': 'Jeevan Wijerathna',
            'title': 'AWS, Azure, and Sitecore Certified Developer with over 10 years of experience in software development. Skilled in .NET technologies, cloud-native architectures, and backend development. Passionate about delivering scalable, maintainable solutions while adhering to best practices like Domain-Driven Design, Clean Code, and Test-Driven Development.',
            'email': 'Jeevan90wijerathna@gmail.com',
            'phone': '+64 22 67 33 146',
            'website': 'iamjeevan.com',
            'location': 'Auckland, New Zealand'
        }
        
        self.work_experience = [
            {
                'title': 'Technical Lead',
                'company': 'Datacom, New Zealand',
                'duration': '2023 - Present',
                'responsibilities': [
                    'Designed and reviewed new features with architects and team members for critical business applications',
                    'Implemented scalable backend solutions using .NET, C#, and Azure technologies',
                    'Enhanced application performance and optimized Azure costs by identifying and addressing bottlenecks',
                    'Led backlog grooming, prioritization, and technical discussions with business users',
                    'Modernized legacy applications by transitioning to cloud-native, decoupled architectures'
                ],
                'projects': [
                    {
                        'name': 'Fonterra Inventory Management System',
                        'description': 'Trace It is an inventory management system for Fonterra, New Zealand\'s largest dairy company',
                        'technologies': 'Oracle Apex, PL/SQL',
                        'achievements': [
                            'Implemented a new feature to track samples inventory',
                            'Introduced software engineering best practices to avoid/minimize deployment risks'
                        ]
                    },
                    {
                        'name': 'PGDB Application Portal',
                        'description': 'Developed a portal for managing licenses and certifications for NZ practitioners',
                        'technologies': '.NET, C#, WEB-API, Identity Server, React, Redux, Azure',
                        'achievements': [
                            'Implemented proactive application monitoring to ensure business continuity',
                            'Delivered scalable architectural solutions for critical business functions'
                        ]
                    },
                    {
                        'name': 'Southern Cross Applications',
                        'description': 'Supported and modernized various business-critical applications',
                        'technologies': '.NET, C#, WEB-API, React, Redux, Azure, Angular, Sitecore',
                        'achievements': [
                            'Ensured business continuity by resolving P1/P2 issues promptly',
                            'Migrated legacy applications to microservices architecture'
                        ]
                    }
                ]
            },
            {
                'title': 'Senior Software Engineer',
                'company': 'Kinesso, Malaysia',
                'duration': '2020 - 2022',
                'responsibilities': [
                    'Designed and implemented modernized legacy applications using microservices architecture',
                    'Developed GRPC services for inter-process communication and optimized database performance',
                    'Automated deployments using Jenkins and introduced DevOps practices',
                    'Integrated various data sources like Snowflake, Redshift, and Athena for analytics tools',
                    'Monitored application performance using DataDog and implemented CI/CD pipelines'
                ],
                'projects': [
                    {
                        'name': 'Magnifiq',
                        'description': 'Developed a business intelligence analytics tool with custom ETL and reporting capabilities',
                        'technologies': '.NET, GRPC, SQL Server, AWS, Docker, Jenkins, Snowflake',
                        'achievements': [
                            'Streamlined database development with Flyway for version control',
                            'Reduced manual deployment efforts by automating processes'
                        ]
                    },
                    {
                        'name': 'Apollo ETL Platform',
                        'description': 'Built an ETL orchestration platform using Dagster and developed a dashboard UI',
                        'technologies': '.NET, Web-API, Dagster, Python, AWS, Angular',
                        'achievements': [
                            'Migrated legacy ETL processes to a modern platform',
                            'Improved ETL performance and scalability'
                        ]
                    }
                ]
            },
            {
                'title': 'Software Engineer',
                'company': 'Sitecore, Malaysia',
                'duration': '2017 - 2022',
                'responsibilities': [
                    'Developed Sitecore platform modules to enhance CMS functionality with Azure Storage and CDN support',
                    'Implemented backend services using WEB-API and C#, and developed front-end applications using Angular',
                    'Set up CI pipelines and automated deployments using PowerShell and TeamCity',
                    'Participated in R&D, sprint planning, and cross-functional team collaboration'
                ]
            },
            {
                'title': 'Software Engineer',
                'company': 'CMS Pvt Ltd, Colombo, Sri Lanka',
                'duration': '2015 - 2017',
                'responsibilities': [
                    'Developed and maintained scalable e-commerce solutions for European markets',
                    'Modernized legacy systems and implemented responsive front-end designs',
                    'Optimized database performance and implemented caching mechanisms'
                ]
            },
            {
                'title': 'Software Engineer',
                'company': 'Bileeta Pvt Ltd, Sri Lanka',
                'duration': 'February 2013 - December 2014',
                'responsibilities': [
                    'Contributed to the development of an award-winning cloud ERP solution',
                    'Designed and implemented backend services and real-time features',
                    'Optimized database performance and implemented stored procedures'
                ]
            }
        ]
        
        self.skills = {
            'Languages & Frameworks': ['C#', '.NET Core', 'ASP.NET', 'React', 'Angular', 'JavaScript', 'TypeScript'],
            'Cloud & Infrastructure': ['Azure', 'AWS', 'Docker', 'Kubernetes', 'Microservices'],
            'DevOps & CI/CD': ['Azure DevOps', 'GitHub Actions', 'Jenkins', 'TeamCity'],
            'Databases & Storage': ['SQL Server', 'PostgreSQL', 'Oracle', 'Azure Storage'],
            'Testing & Quality': ['XUnit', 'SpecFlow', 'Pester', 'TDD', 'Clean Code'],
            'Architecture & Design': ['Domain-Driven Design', 'Clean Architecture', 'SOLID Principles']
        }
        
        self.education = {
            'institution': 'National School of Business Management, Sri Lanka',
            'degree': 'Bachelor of Science in Management Information Systems',
            'year': '2014'
        }
        
        self.certifications = [
            'AWS Certified Solution Architect Associate',
            'AWS Certified Developer',
            'Azure Certified Developer',
            'Sitecore 9 Certified Developer',
            'SAFe 4.0 Practitioner'
        ]
        
        self.personal_projects = [
            {
                'name': 'Resume Pro',
                'url': 'resumepro.iamjeevan.com',
                'description': 'Built an ATS-friendly resume generator using Next.js, React, and Tailwind CSS. Implemented modern UI with real-time preview and multiple templates.',
                'technologies': 'TypeScript, Next.js, React, Tailwind CSS'
            },
            {
                'name': 'Vanish Notes',
                'url': 'vanishnotes.iamjeevan.com',
                'description': 'Developed a self-destructive note sharing application. Implemented end-to-end encryption and auto-deletion features.',
                'technologies': 'Next.js, Node.js, MongoDB, Crypto.js'
            },
            {
                'name': 'NZ Salary Calculator',
                'url': 'nzsalarycalculator.iamjeevan.com',
                'description': 'Created a comprehensive salary calculator for New Zealand employees. Includes tax calculations, KiwiSaver, and other deductions.',
                'technologies': 'React, TypeScript, Tailwind CSS'
            },
            {
                'name': 'Flash Card Fest',
                'url': 'flash-card-fest.netlify.app',
                'description': 'Designed an AI-powered study flash card generator. Integrated with OpenAI API for automated content generation.',
                'technologies': 'React, OpenAI API, Node.js, Express'
            }
        ]


class ATSFriendlyWordGenerator:
    """Generates ATS-friendly Word documents with professional formatting"""
    
    def __init__(self, output_path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_document()
        self._create_styles()
        
    def _setup_document(self):
        """Setup document margins and basic formatting"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def _create_styles(self):
        """Create ATS-friendly styles"""
        styles = self.doc.styles
        
        # Name style
        name_style = styles.add_style('ATSName', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Arial'
        name_style.font.size = Pt(18)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0, 0, 0)
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_style.paragraph_format.space_after = Pt(6)
        
        # Contact info style
        contact_style = styles.add_style('ATSContact', WD_STYLE_TYPE.PARAGRAPH)
        contact_style.font.name = 'Arial'
        contact_style.font.size = Pt(10)
        contact_style.font.color.rgb = RGBColor(0, 0, 0)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style.paragraph_format.space_after = Pt(12)
        
        # Professional title style
        title_style = styles.add_style('ATSTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(11)
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(16)
        title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Section header style
        section_style = styles.add_style('ATSSectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Arial'
        section_style.font.size = Pt(12)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 0, 0)
        section_style.paragraph_format.space_before = Pt(16)
        section_style.paragraph_format.space_after = Pt(8)
        
        # Job title style
        job_title_style = styles.add_style('ATSJobTitle', WD_STYLE_TYPE.PARAGRAPH)
        job_title_style.font.name = 'Arial'
        job_title_style.font.size = Pt(11)
        job_title_style.font.bold = True
        job_title_style.font.color.rgb = RGBColor(0, 0, 0)
        job_title_style.paragraph_format.space_after = Pt(2)
        
        # Company style
        company_style = styles.add_style('ATSCompany', WD_STYLE_TYPE.PARAGRAPH)
        company_style.font.name = 'Arial'
        company_style.font.size = Pt(10)
        company_style.font.italic = True
        company_style.font.color.rgb = RGBColor(0, 0, 0)
        company_style.paragraph_format.space_after = Pt(2)
        
        # Duration style
        duration_style = styles.add_style('ATSDuration', WD_STYLE_TYPE.PARAGRAPH)
        duration_style.font.name = 'Arial'
        duration_style.font.size = Pt(10)
        duration_style.font.color.rgb = RGBColor(0, 0, 0)
        duration_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        body_style = styles.add_style('ATSBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(10)
        body_style.font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.space_after = Pt(4)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Bullet style
        bullet_style = styles.add_style('ATSBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Arial'
        bullet_style.font.size = Pt(10)
        bullet_style.font.color.rgb = RGBColor(0, 0, 0)
        bullet_style.paragraph_format.space_after = Pt(3)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def add_header(self, data):
        """Add professional header section"""
        # Name
        name_para = self.doc.add_paragraph(data.personal_info['name'], style='ATSName')
        
        # Contact info
        contact_info = f"{data.personal_info['email']} | {data.personal_info['phone']} | {data.personal_info['website']} | {data.personal_info['location']}"
        contact_para = self.doc.add_paragraph(contact_info, style='ATSContact')
        
        # Professional summary
        title_para = self.doc.add_paragraph(data.personal_info['title'], style='ATSTitle')
        
        # Add separator line
        separator = self.doc.add_paragraph()
        separator.paragraph_format.space_after = Pt(6)
        
    def add_section_header(self, title):
        """Add section header with consistent formatting"""
        # Add underline to section headers
        header_para = self.doc.add_paragraph(title.upper(), style='ATSSectionHeader')
        # Add border bottom
        p = header_para._element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:sectPrChange')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        
    def add_work_experience(self, data):
        """Add work experience section"""
        self.add_section_header("Professional Experience")
        
        for job in data.work_experience:
            # Job title and duration in the same paragraph
            job_header = self.doc.add_paragraph(style='ATSJobTitle')
            job_header.add_run(job['title']).bold = True
            
            # Add tabs to right-align duration
            tab_stops = job_header.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            job_header.add_run(f"\t{job['duration']}")
            job_header.runs[1].bold = False
            job_header.runs[1].font.size = Pt(10)
            
            # Company
            company_para = self.doc.add_paragraph(job['company'], style='ATSCompany')
            
            # Responsibilities
            for responsibility in job['responsibilities']:
                bullet_para = self.doc.add_paragraph(f"• {responsibility}", style='ATSBullet')
            
            # Projects (if available)
            if 'projects' in job and job['projects']:
                self.doc.add_paragraph()
                projects_header = self.doc.add_paragraph("Key Projects:", style='ATSBody')
                projects_header.runs[0].bold = True
                
                for project in job['projects']:
                    project_para = self.doc.add_paragraph(style='ATSBody')
                    project_para.add_run(f"{project['name']}: ").bold = True
                    project_para.add_run(project['description'])
                    
                    tech_para = self.doc.add_paragraph(style='ATSBody')
                    tech_para.add_run("Technologies: ").italic = True
                    tech_para.add_run(project['technologies'])
                    
                    if 'achievements' in project:
                        for achievement in project['achievements']:
                            achievement_para = self.doc.add_paragraph(f"  • {achievement}", style='ATSBullet')
            
            # Add spacing between jobs
            self.doc.add_paragraph()
    
    def add_skills(self, data):
        """Add skills section in a clean format"""
        self.add_section_header("Technical Skills")
        
        for category, skills in data.skills.items():
            skills_para = self.doc.add_paragraph(style='ATSBody')
            skills_para.add_run(f"{category}: ").bold = True
            skills_para.add_run(', '.join(skills))
        
        self.doc.add_paragraph()
    
    def add_education(self, data):
        """Add education section"""
        self.add_section_header("Education")
        
        # Degree and year in the same paragraph
        education_header = self.doc.add_paragraph(style='ATSBody')
        education_header.add_run(data.education['degree']).bold = True
        
        # Add tabs to right-align year
        tab_stops = education_header.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        education_header.add_run(f"\t{data.education['year']}")
        
        # Institution
        institution_para = self.doc.add_paragraph(data.education['institution'], style='ATSBody')
        self.doc.add_paragraph()
    
    def add_certifications(self, data):
        """Add certifications section"""
        self.add_section_header("Certifications")
        
        for cert in data.certifications:
            cert_para = self.doc.add_paragraph(f"• {cert}", style='ATSBullet')
        
        self.doc.add_paragraph()
    
    def add_personal_projects(self, data):
        """Add personal projects section"""
        self.add_section_header("Personal Projects")
        
        for project in data.personal_projects:
            project_para = self.doc.add_paragraph(style='ATSBody')
            project_para.add_run(f"{project['name']} ").bold = True
            project_para.add_run(f"({project['url']})")
            
            desc_para = self.doc.add_paragraph(project['description'], style='ATSBody')
            
            tech_para = self.doc.add_paragraph(style='ATSBody')
            tech_para.add_run("Technologies: ").italic = True
            tech_para.add_run(project['technologies'])
            
            self.doc.add_paragraph()
    
    def generate(self, data):
        """Generate the complete Word document"""
        self.add_header(data)
        self.add_work_experience(data)
        self.add_skills(data)
        self.add_education(data)
        self.add_certifications(data)
        self.add_personal_projects(data)
        
        # Save the document
        self.doc.save(self.output_path)
        print(f"✅ Professional ATS-friendly Word document generated: {self.output_path}")


def main():
    parser = argparse.ArgumentParser(description='Generate professional ATS-friendly Word resume')
    parser.add_argument('--output', '-o',
                       default=None,
                       help='Output Word file path (default: auto-generated)')
    
    args = parser.parse_args()
    
    # Get the script directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    
    # Output file path
    if args.output:
        output_path = Path(args.output)
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"Jeevan_Wijerathna_Resume_Professional_{timestamp}.docx"
        output_path = project_root / output_filename
    
    print(f"🔄 Generating professional Word resume...")
    print(f"📄 Output: {output_path}")
    
    # Create resume data
    resume_data = ResumeData()
    
    # Generate Word document
    word_generator = ATSFriendlyWordGenerator(str(output_path))
    word_generator.generate(resume_data)
    
    # Show file info
    if output_path.exists():
        file_size = output_path.stat().st_size / 1024
        print(f"📊 File size: {file_size:.1f} KB")
        print(f"🎉 Professional Word resume generated successfully!")
        print(f"\n✅ ATS Optimization Features:")
        print(f"   • Standard fonts (Arial)")
        print(f"   • Clean, structured layout")
        print(f"   • Proper text hierarchy")
        print(f"   • No complex formatting")
        print(f"   • Fully editable text")
        print(f"   • Standard document structure")
    else:
        print("❌ Failed to generate Word document")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())